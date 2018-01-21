def export_docx(y):
    from docx import Document
    import time

    rows = len(y[1])
    document = Document()

    document.add_heading('旋风分离器计算书', 0)
    p = document.add_paragraph('免责声明：本计算书内容由热能小助手网站提供，其数据仅供参考，请用户自行校核。')

    document.add_heading('计算结果表', level=1)

    table = document.add_table(rows=rows+1, cols=4)
    table.style = 'Table Grid'
    table.autofit = True
    table.cell(0, 0).text = "名称"
    table.cell(0, 1).text = "数值"
    table.cell(0, 2).text = "单位"
    table.cell(0, 3).text = "备注"
    i=1
    for item in y[1]:
        table.cell(i,0).text = item['text']
        table.cell(i,1).text = item['value']
        table.cell(i,2).text = item['unit']
        table.cell(i,3).text = item['memo']
        i += 1

    document.add_page_break()
    trick = int(time.time()).__str__()
    fn = 'result'+trick+'.docx'
    document.save('static/results/'+fn)
    return fn


# 字典拆分成变量
def mody_01_test(x):
    for key, val in x.items():
        exec(key + '=val')
    tmp = 2
    print(locals()['T'])
    return "OK"